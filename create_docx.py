from docx import Document

# 문서 생성
doc = Document()

# 섹션 4: 사업 운영 계획
doc.add_heading('4. 사업 운영 계획', level=2)

# 주요 추진 일정
doc.add_paragraph('**주요 추진 일정:**')
table_1 = doc.add_table(rows=4, cols=2)
data_1 = [
    ['기간', '주요 활동'],
    ['2025년 1월', '공장 설계 및 부지 정비'],
    ['2025년 3월', '설비 설치 및 시험 가동'],
    ['2025년 5월', '상업 생산 개시']
]
for i, row in enumerate(table_1.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data_1[i][j]

# 운영 인력
doc.add_heading('운영 인력', level=3)
table_2 = doc.add_table(rows=2, cols=2)
data_2 = [
    ['구분', '인원수'],
    ['현지 채용 인원', '20명 이상']
]
for i, row in enumerate(table_2.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data_2[i][j]

# 파일 저장
file_path = "사업운영계획.docx"
doc.save(file_path)
print(f"DOCX 파일이 생성되었습니다: {file_path}")
